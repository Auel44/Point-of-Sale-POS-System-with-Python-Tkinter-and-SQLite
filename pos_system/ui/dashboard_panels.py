"""Embedded dashboard panels for cashier workflow modules."""

from __future__ import annotations

import csv
from datetime import datetime
import os
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
from typing import Any

from modules.auth import change_password
from modules.customers import (
	add_customer,
	get_all_customers,
	get_purchase_history,
	search_customers,
	update_customer,
)
from modules.audit import list_recent as list_recent_audit_logs
from modules.inventory import adjust_stock, get_inventory_log
from modules.permissions import has_permission
from modules.payments import list_payments_by_date_range, list_recent_payments
from modules.products import (
	add_product,
	delete_product,
	get_all_products,
	get_product_by_barcode,
	get_product_by_id,
	search_products,
	update_product,
)
from modules.receipts import generate_receipt
from modules.sales import get_sale_by_id, get_sales_by_date_range
from modules.tickets import (
	close_ticket,
	list_all_tickets,
	resolve_ticket,
)
from modules.users import (
	create_user,
	delete_user,
	get_all_users,
	is_password_in_use,
	reset_user_password,
	update_user_details,
	update_user_role,
)
from modules.reports import (
	cashier_report,
	daily_sales_report,
	inventory_report,
	product_performance_report,
	weekly_sales_report,
)
from ui.payment_screen import PaymentScreen
from ui.theme import apply_modern_theme, themed_confirm_dialog
from utils.validators import (
	generate_password,
	generate_username_from_fullname,
	is_non_empty,
	is_valid_email,
	is_valid_price,
	is_valid_quantity,
	password_policy_error,
)

try:
	from PIL import Image, ImageDraw, ImageFont  # type: ignore[import-not-found]
	PIL_AVAILABLE = True
except Exception:
	PIL_AVAILABLE = False

try:
	from docx import Document  # type: ignore[import-not-found]
	DOCX_AVAILABLE = True
except Exception:
	DOCX_AVAILABLE = False


def _save_credentials_docx_to_desktop(
	username: str,
	password: str,
	title: str,
	full_name: str = "",
	role: str = "",
) -> str:
	"""Save username/password details as a .docx file on the current user's desktop."""
	if not DOCX_AVAILABLE:
		raise RuntimeError("DOCX export is unavailable. Install python-docx first.")

	desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
	os.makedirs(desktop_path, exist_ok=True)
	timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
	filename = f"credentials_{username}_{timestamp}.docx"
	output_path = os.path.join(desktop_path, filename)

	doc = Document()
	doc.add_heading(title, level=1)
	doc.add_paragraph(f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
	doc.add_paragraph(f"Username: {username}")
	doc.add_paragraph(f"Password: {password}")
	if full_name:
		doc.add_paragraph(f"Full Name: {full_name}")
	if role:
		doc.add_paragraph(f"Role: {role}")
	doc.add_paragraph("Share this document securely.")
	doc.save(output_path)
	return output_path


def _show_confirmation_dialog(parent: tk.Misc, title: str, message: str) -> bool:
	"""Render a themed confirmation dialog and return True when user confirms."""
	return themed_confirm_dialog(parent, title, message)


class SalesPanel(ttk.Frame):
	"""In-dashboard sales cart with checkout and daily sales history."""

	def __init__(self, parent: tk.Misc, user: dict[str, Any] | None = None) -> None:
		super().__init__(parent, style="Card.TFrame", padding=(12, 10))
		self._user = user or {}
		self._role = str(self._user.get("role", "Cashier"))
		self._is_admin = self._role == "Admin"
		self._is_cashier = self._role == "Cashier"
		self._is_manager = self._role == "Manager"
		self._cart: list[dict[str, Any]] = []
		self._query_var = tk.StringVar()
		self._quantity_var = tk.StringVar(value="1")
		self._discount_var = tk.StringVar(value="0")
		today = datetime.now().strftime("%Y-%m-%d")
		self._history_start_var = tk.StringVar(value=today)
		self._history_end_var = tk.StringVar(value=today)
		self._cashier_filter_var = tk.StringVar(value="All")
		self._history_rows: list[dict[str, Any]] = []
		self._search_results: list[dict[str, Any]] = []
		self._error_var = tk.StringVar()
		self._build_ui()
		if not self._is_manager and not self._is_admin:
			self._refresh_totals()
		self._refresh_history()
		self.after(5000, self._auto_refresh_history)

	def _build_ui(self) -> None:
		nb = ttk.Notebook(self)
		nb.pack(expand=True, fill="both")

		cart_tab = ttk.Frame(nb, padding=6, style="Card.TFrame")
		history_tab = ttk.Frame(nb, padding=6, style="Card.TFrame")
		if not self._is_admin and not self._is_manager:
			nb.add(cart_tab, text="Cart")
		nb.add(history_tab, text="Sales History")

		top = ttk.Frame(cart_tab, style="Card.TFrame")
		top.pack(fill="x", pady=(0, 8))

		ttk.Label(top, text="Search Product:", style="SectionSub.TLabel").pack(side="left", padx=(0, 6))
		self._search_combo = ttk.Combobox(top, textvariable=self._query_var, width=40)
		self._search_combo.pack(side="left")
		self._search_combo.bind("<KeyRelease>", self._on_search_key)
		self._search_combo.bind("<<ComboboxSelected>>", self._on_search_select)
		self._search_combo.bind("<Return>", lambda _e: self._add_item())

		ttk.Label(top, text="Qty:", style="SectionSub.TLabel").pack(side="left", padx=(10, 6))
		qty_entry = ttk.Entry(top, textvariable=self._quantity_var, width=6, justify="right")
		qty_entry.pack(side="left")
		qty_entry.bind("<Return>", lambda _e: self._add_item())

		ttk.Button(top, text="Add Item", command=self._add_item, style="Primary.TButton").pack(side="left", padx=(8, 0))
		ttk.Button(top, text="Make Payment", command=self._checkout, style="Primary.TButton").pack(side="left", padx=(8, 0))

		cols = ("product_id", "product_name", "unit_price", "quantity", "subtotal")
		self._tree = ttk.Treeview(cart_tab, columns=cols, show="headings", height=11, selectmode="browse")
		headers = {
			"product_id": ("ID", 60, "center"),
			"product_name": ("Product", 280, "w"),
			"unit_price": ("Price", 100, "e"),
			"quantity": ("Qty", 70, "center"),
			"subtotal": ("Subtotal", 120, "e"),
		}
		for col, (label, width, anchor) in headers.items():
			self._tree.heading(col, text=label)
			self._tree.column(col, width=width, anchor=anchor)
		self._tree.pack(fill="both", expand=True)

		actions = ttk.Frame(cart_tab, style="Card.TFrame")
		actions.pack(fill="x", pady=(8, 6))
		ttk.Button(actions, text="Edit Selected Qty", command=self._edit_selected_quantity).pack(side="left", padx=(0, 6))
		ttk.Button(actions, text="Remove", command=self._remove_item, style="Danger.TButton").pack(side="left")

		totals = ttk.Frame(cart_tab, style="Card.TFrame")
		totals.pack(fill="x")
		ttk.Label(totals, text="Discount (GHc):", style="SectionSub.TLabel").pack(side="left", padx=(0, 6))
		discount_entry = ttk.Entry(totals, textvariable=self._discount_var, width=8, justify="right")
		discount_entry.pack(side="left")
		discount_entry.bind("<KeyRelease>", lambda _e: self._refresh_totals())

		self._summary_lbl = ttk.Label(totals, text="Subtotal: GHc 0.00  |  Total: GHc 0.00", style="InfoBadge.TLabel")
		self._summary_lbl.pack(side="right")

		ttk.Label(cart_tab, textvariable=self._error_var, foreground="red").pack(anchor="w", pady=(6, 0))

		hf = ttk.Frame(history_tab, style="Card.TFrame")
		hf.pack(fill="x", pady=(0, 8))
		ttk.Label(hf, text="Start:", style="SectionSub.TLabel").pack(side="left", padx=(0, 6))
		start_entry = ttk.Entry(hf, textvariable=self._history_start_var, width=12)
		start_entry.pack(side="left")
		start_entry.bind("<Return>", lambda _e: self._refresh_history())
		ttk.Label(hf, text="End:", style="SectionSub.TLabel").pack(side="left", padx=(8, 6))
		end_entry = ttk.Entry(hf, textvariable=self._history_end_var, width=12)
		end_entry.pack(side="left")
		end_entry.bind("<Return>", lambda _e: self._refresh_history())
		ttk.Button(hf, text="Today", command=self._set_today_filter).pack(side="left", padx=(6, 0))
		ttk.Button(hf, text="Refresh", command=self._refresh_history, style="Primary.TButton").pack(side="left", padx=(6, 0))
		if not self._is_cashier:
			ttk.Label(hf, text="Cashier:", style="SectionSub.TLabel").pack(side="left", padx=(8, 6))
			self._cashier_filter_combo = ttk.Combobox(
				hf,
				textvariable=self._cashier_filter_var,
				state="readonly",
				width=16,
			)
			self._cashier_filter_combo.pack(side="left")
			self._cashier_filter_combo.configure(values=["All"])
			self._cashier_filter_var.set("All")
			self._cashier_filter_combo.bind("<<ComboboxSelected>>", lambda _e: self._refresh_history())
			ttk.Button(hf, text="Export CSV", command=self._export_history_csv).pack(side="left", padx=(6, 0))

		h_cols = (
			("sale_id", "date", "cashier_username", "payment_method", "total_amount", "today_sale")
			if not self._is_cashier
			else ("sale_id", "date", "payment_method", "total_amount", "today_sale")
		)
		self._history_tree = ttk.Treeview(history_tab, columns=h_cols, show="headings", height=14)
		headers = (
			{
				"sale_id": ("Sale #", 80, "center"),
				"date": ("Date/Time", 190, "w"),
				"cashier_username": ("Cashier", 140, "w"),
				"payment_method": ("Method", 120, "w"),
				"total_amount": ("Total", 100, "e"),
				"today_sale": ("Today?", 80, "center"),
			}
			if not self._is_cashier
			else {
				"sale_id": ("Sale #", 80, "center"),
				"date": ("Date/Time", 210, "w"),
				"payment_method": ("Method", 130, "w"),
				"total_amount": ("Total", 110, "e"),
				"today_sale": ("Today?", 80, "center"),
			}
		)
		for col, (label, width, anchor) in headers.items():
			self._history_tree.heading(col, text=label)
			self._history_tree.column(col, width=width, anchor=anchor)
		self._history_tree.pack(fill="both", expand=True)

		self._history_status_var = tk.StringVar(value="0 sale(s)")
		ttk.Label(history_tab, textvariable=self._history_status_var, style="InfoBadge.TLabel").pack(anchor="w", pady=(8, 0))

	def _set_today_filter(self) -> None:
		today = datetime.now().strftime("%Y-%m-%d")
		self._history_start_var.set(today)
		self._history_end_var.set(today)
		self._refresh_history()

	def _refresh_history(self) -> None:
		for row in self._history_tree.get_children():
			self._history_tree.delete(row)
		start = self._history_start_var.get().strip()
		end = self._history_end_var.get().strip()
		try:
			start_date = datetime.strptime(start, "%Y-%m-%d")
			end_date = datetime.strptime(end, "%Y-%m-%d")
		except ValueError:
			self._history_rows = []
			self._history_status_var.set("Invalid date format. Use YYYY-MM-DD.")
			return
		if start_date > end_date:
			self._history_rows = []
			self._history_status_var.set("Start date cannot be after end date.")
			return
		if not self._is_cashier:
			self._refresh_cashier_filter_options()

		cashier_filter = None
		if not self._is_cashier:
			selected_cashier = self._cashier_filter_var.get().strip()
			if selected_cashier and selected_cashier != "All":
				cashier_filter = selected_cashier

		rows = get_sales_by_date_range(
			start,
			end,
			user_id=int(self._user["user_id"]) if self._is_cashier and "user_id" in self._user else None,
			cashier_username=cashier_filter,
		)
		self._history_rows = rows

		today_key = datetime.now().strftime("%Y-%m-%d")
		today_count = 0
		today_total = 0.0
		period_total = 0.0
		for s in rows:
			is_today = str(s["date"]).startswith(today_key)
			amount = float(s["total_amount"])
			period_total += amount
			if is_today:
				today_count += 1
				today_total += amount
			if self._is_cashier:
				values = (
					s["sale_id"],
					s["date"],
					s.get("payment_method", ""),
					f"{amount:.2f}",
					"Yes" if is_today else "No",
				)
			else:
				values = (
					s["sale_id"],
					s["date"],
					s.get("cashier_username", ""),
					s.get("payment_method", ""),
					f"{amount:.2f}",
					"Yes" if is_today else "No",
				)
			self._history_tree.insert("", "end", values=values)
		if self._is_cashier:
			self._history_status_var.set(
				f"{len(rows)} sale(s) from {start} to {end} | Today in selected period: {today_count} sale(s), GHc {today_total:.2f}"
			)
		else:
			cashier_name = self._cashier_filter_var.get().strip()
			if cashier_name and cashier_name != "All":
				self._history_status_var.set(
					f"{len(rows)} sale(s) for {cashier_name} from {start} to {end} | Total by cashier: GHc {period_total:.2f}"
				)
			else:
				self._history_status_var.set(f"{len(rows)} sale(s) from {start} to {end} | Total: GHc {period_total:.2f}")

	def _refresh_cashier_filter_options(self) -> None:
		cashier_names = sorted(
			{
				str(u.get("username", "")).strip()
				for u in get_all_users()
				if str(u.get("role", "")).strip() == "Cashier" and str(u.get("username", "")).strip()
			}
		)
		current = self._cashier_filter_var.get().strip() or "All"
		self._cashier_filter_combo.configure(values=["All"] + cashier_names)
		if current != "All" and current not in cashier_names:
			self._cashier_filter_var.set("All")
		elif current:
			self._cashier_filter_var.set(current)

	def _export_history_csv(self) -> None:
		if self._is_cashier:
			self._history_status_var.set("Cashier accounts cannot export sales records.")
			return
		if not self._history_rows:
			self._history_status_var.set("No sales rows to export for the current filter.")
			return
		path = filedialog.asksaveasfilename(
			parent=self,
			defaultextension=".csv",
			filetypes=[("CSV Files", "*.csv")],
			initialfile="sales_history.csv",
		)
		if not path:
			return
		try:
			with open(path, "w", newline="", encoding="utf-8") as csv_file:
				fieldnames = ["sale_id", "date", "cashier_username", "payment_method", "total_amount"]
				writer = csv.DictWriter(csv_file, fieldnames=fieldnames)
				writer.writeheader()
				for row in self._history_rows:
					writer.writerow(
						{
							"sale_id": row["sale_id"],
							"date": row["date"],
							"cashier_username": row.get("cashier_username", ""),
							"payment_method": row.get("payment_method", ""),
							"total_amount": f"{float(row['total_amount']):.2f}",
						}
					)
			self._history_status_var.set("Sales history exported successfully.")
		except OSError as exc:
			self._history_status_var.set(f"Export failed: {exc}")

	def _auto_refresh_history(self) -> None:
		try:
			self._refresh_history()
		except Exception:
			pass
		if self.winfo_exists():
			self.after(5000, self._auto_refresh_history)

	def _checkout(self) -> None:
		self._error_var.set("")
		if not self._cart:
			self._error_var.set("Cart is empty.")
			return
		if "user_id" not in self._user:
			self._error_var.set("User session missing. Please log in again.")
			return
		try:
			discount = max(0.0, float(self._discount_var.get().strip() or "0"))
		except ValueError:
			self._error_var.set("Invalid discount value.")
			return
		subtotal = sum(float(i["unit_price"]) * int(i["quantity"]) for i in self._cart)
		grand_total = max(0.0, subtotal - discount)
		cart_snapshot = [
			{
				"product_id": item["product_id"],
				"product_name": item["product_name"],
				"quantity": item["quantity"],
				"price": item["unit_price"],
			}
			for item in self._cart
		]
		PaymentScreen(
			self,
			user=self._user,
			cart=cart_snapshot,
			discount=discount,
			tax=0.0,
			grand_total=grand_total,
			on_success=self._on_sale_completed,
		)

	def _on_sale_completed(self, _sale_id: int) -> None:
		"""Clear cart and refresh history after a successful payment."""
		self._cart.clear()
		self._query_var.set("")
		self._quantity_var.set("1")
		self._discount_var.set("0")
		self._refresh_rows()
		self._refresh_history()

	def _on_search_key(self, _event: Any = None) -> None:
		query = self._query_var.get().strip()
		if not query:
			self._search_results = []
			self._search_combo.configure(values=[])
			return
		matches = search_products(query)[:15]
		self._search_results = matches
		values = [
			f"{p['product_name']} (ID:{p['product_id']}) | {p.get('barcode', '')}"
			for p in matches
		]
		self._search_combo.configure(values=values)
		if values:
			try:
				self._search_combo.event_generate("<Down>")
			except Exception:
				pass

	def _on_search_select(self, _event: Any = None) -> None:
		# Keep selected value as-is; parsing happens in _resolve_product
		return

	def _resolve_product(self, query: str) -> dict[str, Any] | None:
		if "ID:" in query:
			try:
				product_id = int(query.split("ID:")[1].split(")")[0])
				product = get_product_by_id(product_id)
				if product is not None:
					return product
			except ValueError:
				pass

		product = get_product_by_barcode(query)
		if product is not None:
			return product
		if query.isdigit():
			product = get_product_by_id(int(query))
			if product is not None:
				return product
		matches = search_products(query)
		if len(matches) == 1:
			return matches[0]
		if len(matches) > 1:
			# When user typed text without selecting a dropdown entry, use first filtered match.
			return matches[0]
		return None

	def _add_item(self) -> None:
		self._error_var.set("")
		query = self._query_var.get().strip()
		if not query:
			self._error_var.set("Enter a barcode, ID, or product name.")
			return
		try:
			qty_to_add = int(self._quantity_var.get().strip())
		except ValueError:
			self._error_var.set("Quantity must be a whole number.")
			return
		if qty_to_add <= 0:
			self._error_var.set("Quantity must be greater than zero.")
			return

		product = self._resolve_product(query)
		if product is None:
			if not self._error_var.get():
				self._error_var.set(f"No product found for '{query}'.")
			return
		if int(product["quantity"]) <= 0:
			self._error_var.set(f"{product['product_name']} is out of stock.")
			return
		for item in self._cart:
			if item["product_id"] == product["product_id"]:
				if int(item["quantity"]) + qty_to_add > int(product["quantity"]):
					msg = f"Only {product['quantity']} unit(s) of {product['product_name']} are available in stock."
					self._error_var.set(msg)
					messagebox.showwarning("Stock Limit", msg, parent=self)
					return
				item["quantity"] += qty_to_add
				self._refresh_rows()
				self._quantity_var.set("1")
				return
		if qty_to_add > int(product["quantity"]):
			msg = f"Only {product['quantity']} unit(s) of {product['product_name']} are available in stock."
			self._error_var.set(msg)
			messagebox.showwarning("Stock Limit", msg, parent=self)
			return
		self._cart.append(
			{
				"product_id": product["product_id"],
				"product_name": product["product_name"],
				"unit_price": float(product["price"]),
				"quantity": qty_to_add,
			}
		)
		self._query_var.set("")
		self._quantity_var.set("1")
		self._refresh_rows()

	def _selected_index(self) -> int | None:
		selected = self._tree.selection()
		if not selected:
			return None
		return self._tree.index(selected[0])

	def _edit_selected_quantity(self) -> None:
		idx = self._selected_index()
		if idx is None:
			self._error_var.set("Select a cart row to edit quantity.")
			return
		item = self._cart[idx]
		new_qty = simpledialog.askinteger(
			"Edit Quantity",
			f"Enter quantity for {item['product_name']}:",
			parent=self,
			minvalue=1,
			initialvalue=int(item["quantity"]),
		)
		if new_qty is None:
			return
		product = get_product_by_id(int(item["product_id"]))
		if product and int(new_qty) > int(product["quantity"]):
			msg = f"Only {product['quantity']} unit(s) of {product['product_name']} are available in stock."
			self._error_var.set(msg)
			messagebox.showwarning("Stock Limit", msg, parent=self)
			return
		item["quantity"] = int(new_qty)
		self._refresh_rows()

	def _remove_item(self) -> None:
		idx = self._selected_index()
		if idx is None:
			return
		self._cart.pop(idx)
		self._refresh_rows()

	def _refresh_rows(self) -> None:
		for row in self._tree.get_children():
			self._tree.delete(row)
		for item in self._cart:
			subtotal = float(item["unit_price"]) * int(item["quantity"])
			self._tree.insert(
				"",
				"end",
				values=(
					item["product_id"],
					item["product_name"],
					f"{float(item['unit_price']):.2f}",
					item["quantity"],
					f"{subtotal:.2f}",
				),
			)
		self._refresh_totals()

	def _refresh_totals(self) -> None:
		subtotal = sum(float(i["unit_price"]) * int(i["quantity"]) for i in self._cart)
		try:
			discount = max(0.0, float(self._discount_var.get().strip() or "0"))
		except ValueError:
			discount = 0.0
		total = max(0.0, subtotal - discount)
		self._summary_lbl.config(text=f"Subtotal: GHc {subtotal:.2f}  |  Total: GHc {total:.2f}")


class PaymentsPanel(ttk.Frame):
	"""In-dashboard payments history panel."""

	def __init__(self, parent: tk.Misc, user: dict[str, Any] | None = None) -> None:
		super().__init__(parent, style="Card.TFrame", padding=(12, 10))
		self._user = user or {}
		self._role = str(self._user.get("role", "Cashier"))
		self._is_cashier = self._role == "Cashier"
		self._is_manager = self._role == "Manager"
		today = datetime.now().strftime("%Y-%m-%d")
		self._start_var = tk.StringVar(value=today)
		self._end_var = tk.StringVar(value=today)
		self._search_var = tk.StringVar()
		self._rows: list[dict[str, Any]] = []
		self._all_rows: list[dict[str, Any]] = []
		self._build_ui()
		self.refresh()

	def _build_ui(self) -> None:
		header = ttk.Frame(self, style="Card.TFrame")
		header.pack(fill="x", pady=(0, 8))
		ttk.Label(header, text="Payment Records", style="SectionTitle.TLabel").pack(side="left")
		self._count_lbl = ttk.Label(header, text="0 records", style="InfoBadge.TLabel")
		self._count_lbl.pack(side="right")
		if self._is_manager:
			ttk.Label(header, text="Start:", style="SectionSub.TLabel").pack(side="left", padx=(14, 6))
			ttk.Entry(header, textvariable=self._start_var, width=12).pack(side="left")
			ttk.Label(header, text="End:", style="SectionSub.TLabel").pack(side="left", padx=(8, 6))
			ttk.Entry(header, textvariable=self._end_var, width=12).pack(side="left")
			ttk.Button(header, text="Apply Range", command=self.refresh, style="Primary.TButton").pack(side="left", padx=(8, 0))
		ttk.Label(header, text="Search:", style="SectionSub.TLabel").pack(side="left", padx=(12, 6))
		search_entry = ttk.Entry(header, textvariable=self._search_var, width=24)
		search_entry.pack(side="left")
		search_entry.bind("<KeyRelease>", lambda _e: self._apply_search_filter())
		if not self._is_cashier:
			ttk.Button(header, text="Export CSV", command=self._export_csv).pack(side="right", padx=(8, 8))
		ttk.Button(header, text="Refresh", command=self.refresh, style="Primary.TButton").pack(
			side="right", padx=(0, 8)
		)

		cols = ("payment_id", "sale_id", "date", "cashier", "method", "sale_total", "amount_paid", "change")
		self._tree = ttk.Treeview(self, columns=cols, show="headings", height=15)
		headers = {
			"payment_id": ("Payment #", 90, "center"),
			"sale_id": ("Sale #", 80, "center"),
			"date": ("Date", 170, "w"),
			"cashier": ("Cashier", 120, "w"),
			"method": ("Method", 120, "w"),
			"sale_total": ("Sale Total", 100, "e"),
			"amount_paid": ("Amount Paid", 110, "e"),
			"change": ("Change", 90, "e"),
		}
		for col, (label, width, anchor) in headers.items():
			self._tree.heading(col, text=label)
			self._tree.column(col, width=width, anchor=anchor)
		self._tree.pack(expand=True, fill="both")

		# Keep all critical actions in the header so they remain visible on resize.

	def refresh(self) -> None:
		if self._is_manager:
			start = self._start_var.get().strip()
			end = self._end_var.get().strip()
			try:
				datetime.strptime(start, "%Y-%m-%d")
				datetime.strptime(end, "%Y-%m-%d")
			except ValueError:
				self._rows = []
				self._count_lbl.config(text="Invalid date format (YYYY-MM-DD)")
				return
			rows = list_payments_by_date_range(start, end)
		else:
			rows = list_recent_payments(
				limit=500,
				user_id=int(self._user["user_id"]) if self._is_cashier and "user_id" in self._user else None,
			)
		self._all_rows = rows
		self._apply_search_filter()

	def _apply_search_filter(self) -> None:
		for row in self._tree.get_children():
			self._tree.delete(row)
		query = self._search_var.get().strip().lower()
		if not query:
			self._rows = list(self._all_rows)
		else:
			self._rows = [
				row
				for row in self._all_rows
				if query in str(row.get("payment_id", "")).lower()
				or query in str(row.get("sale_id", "")).lower()
				or query in str(row.get("cashier_username", "")).lower()
			]
		for row in self._rows:
			self._tree.insert(
				"",
				"end",
				values=(
					row["payment_id"],
					row["sale_id"],
					row["date"],
					row.get("cashier_username", ""),
					row["payment_method"],
					f"{float(row['sale_total']):.2f}",
					f"{float(row['amount_paid']):.2f}",
					f"{float(row['change_given']):.2f}",
				),
			)
		total_in_range = sum(float(r.get("sale_total", 0) or 0) for r in self._rows)
		if self._is_manager:
			self._count_lbl.config(text=f"{len(self._rows)} records | Total GHc {total_in_range:.2f}")
		else:
			self._count_lbl.config(text=f"{len(self._rows)} records")

	def _export_csv(self) -> None:
		if self._is_cashier:
			messagebox.showinfo("Export", "Cashier accounts cannot export payment records.", parent=self)
			return
		if not self._rows:
			messagebox.showinfo("Export", "No payment rows to export.", parent=self)
			return
		path = filedialog.asksaveasfilename(
			parent=self,
			defaultextension=".csv",
			filetypes=[("CSV Files", "*.csv")],
			initialfile="payment_records.csv",
		)
		if not path:
			return
		try:
			with open(path, "w", newline="", encoding="utf-8") as csv_file:
				writer = csv.DictWriter(
					csv_file,
					fieldnames=[
						"payment_id",
						"sale_id",
						"date",
						"payment_method",
						"sale_total",
						"amount_paid",
						"change_given",
					],
				)
				writer.writeheader()
				for row in self._rows:
					writer.writerow(row)
			messagebox.showinfo("Export", "Payment records exported.", parent=self)
		except OSError as exc:
			messagebox.showerror("Export Error", str(exc), parent=self)


class ReceiptsPanel(ttk.Frame):
	"""In-dashboard receipt lookup and preview panel."""

	def __init__(self, parent: tk.Misc, user: dict[str, Any] | None = None) -> None:
		super().__init__(parent, style="Card.TFrame", padding=(12, 10))
		self._user = user or {}
		self._is_cashier = str(self._user.get("role", "Cashier")) == "Cashier"
		self._sale_id_var = tk.StringVar()
		self._error_var = tk.StringVar()
		self._current_receipt_text = ""
		self._current_sale_id: int | None = None
		self._build_ui()

	def _build_ui(self) -> None:
		top = ttk.Frame(self, style="Card.TFrame")
		top.pack(fill="x", pady=(0, 8))
		ttk.Label(top, text="Sale ID:", style="SectionSub.TLabel").pack(side="left", padx=(0, 6))
		entry = ttk.Entry(top, textvariable=self._sale_id_var, width=18)
		entry.pack(side="left")
		entry.bind("<Return>", lambda _e: self._load_receipt())
		ttk.Button(top, text="Load Receipt", command=self._load_receipt, style="Primary.TButton").pack(side="left", padx=(6, 0))
		ttk.Button(top, text="Download PNG", command=self._download_receipt).pack(side="left", padx=(6, 0))

		ttk.Label(self, textvariable=self._error_var, foreground="red").pack(anchor="w", pady=(0, 6))

		text_frame = ttk.Frame(self, style="Card.TFrame")
		text_frame.pack(expand=True, fill="both")
		self._text = tk.Text(text_frame, font=("Courier New", 9), wrap="none", state="disabled")
		vsb = ttk.Scrollbar(text_frame, orient="vertical", command=self._text.yview)
		hsb = ttk.Scrollbar(text_frame, orient="horizontal", command=self._text.xview)
		self._text.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
		vsb.pack(side="right", fill="y")
		hsb.pack(side="bottom", fill="x")
		self._text.pack(side="left", fill="both", expand=True)

	def _load_receipt(self) -> None:
		self._error_var.set("")
		raw = self._sale_id_var.get().strip()
		if not raw or not raw.isdigit():
			self._error_var.set("Enter a valid numeric sale ID.")
			return
		sale = get_sale_by_id(int(raw))
		if sale is None:
			self._error_var.set(f"Sale #{raw} was not found.")
			return
		if self._is_cashier and "user_id" in self._user and int(sale.get("user_id", -1)) != int(self._user["user_id"]):
			self._error_var.set("Cashier accounts can only view their own receipts.")
			return
		receipt = generate_receipt(int(raw))
		if not receipt:
			self._error_var.set(f"Sale #{raw} was not found.")
			return
		self._text.configure(state="normal")
		self._text.delete("1.0", "end")
		self._text.insert("1.0", receipt)
		self._text.configure(state="disabled")
		self._current_receipt_text = receipt
		self._current_sale_id = int(raw)

	def _download_receipt(self) -> None:
		if not self._current_receipt_text:
			self._error_var.set("Load a receipt first before downloading.")
			return
		if not PIL_AVAILABLE:
			self._error_var.set("PNG export requires Pillow. Please install pillow.")
			return
		default_name = (
			f"receipt_{self._current_sale_id}.png" if self._current_sale_id is not None else "receipt.png"
		)
		path = filedialog.asksaveasfilename(
			parent=self,
			defaultextension=".png",
			filetypes=[("PNG Image", "*.png")],
			initialfile=default_name,
		)
		if not path:
			return
		try:
			lines = self._current_receipt_text.splitlines() or [""]
			font = ImageFont.load_default()
			line_height = 18
			padding = 16
			width = 860
			height = max(220, padding * 2 + line_height * len(lines))
			image = Image.new("RGB", (width, height), "white")
			draw = ImageDraw.Draw(image)
			y = padding
			for line in lines:
				draw.text((padding, y), line, fill="black", font=font)
				y += line_height
			image.save(path, format="PNG")
			self._error_var.set("Receipt PNG downloaded successfully.")
		except OSError as exc:
			self._error_var.set(f"Download failed: {exc}")


class ProductManagementPanel(ttk.Frame):
	"""In-dashboard product CRUD panel."""

	def __init__(self, parent: tk.Misc, user: dict[str, Any] | None = None) -> None:
		super().__init__(parent, style="Card.TFrame", padding=(12, 10))
		self._user = user or {}
		self._role = str(self._user.get("role", "Cashier"))
		self._is_admin = self._role == "Admin"
		self._search_var = tk.StringVar()
		self._name_var = tk.StringVar()
		self._category_var = tk.StringVar()
		self._price_var = tk.StringVar()
		self._quantity_var = tk.StringVar()
		self._barcode_var = tk.StringVar()
		self._status_var = tk.StringVar()
		self._build_ui()
		self.refresh()

	def _build_ui(self) -> None:
		header = ttk.Frame(self, style="Card.TFrame")
		header.pack(fill="x", pady=(0, 8))
		ttk.Label(header, text="Product Management", style="SectionTitle.TLabel").pack(side="left")

		search_row = ttk.Frame(self, style="Card.TFrame")
		search_row.pack(fill="x", pady=(0, 8))
		ttk.Label(search_row, text="Search:", style="SectionSub.TLabel").pack(side="left", padx=(0, 6))
		entry = ttk.Entry(search_row, textvariable=self._search_var, width=30)
		entry.pack(side="left")
		entry.bind("<KeyRelease>", lambda _e: self.refresh())
		ttk.Button(search_row, text="Clear", command=self._clear_search).pack(side="left", padx=(6, 0))

		if not self._is_admin:
			form_wrap = ttk.Frame(self, style="Card.TFrame", padding=(0, 8, 0, 0))
			form_wrap.pack(fill="x")
			form = ttk.Frame(form_wrap, style="Card.TFrame")
			form.pack(fill="x")
			form.columnconfigure(1, weight=1)
			form.columnconfigure(3, weight=1)
			fields = [
				("Name", self._name_var),
				("Category", self._category_var),
				("Price", self._price_var),
				("Quantity", self._quantity_var),
				("Barcode", self._barcode_var),
			]
			for idx, (label, var) in enumerate(fields):
				row = idx // 2
				col_base = (idx % 2) * 2
				ttk.Label(form, text=f"{label}:", style="SectionSub.TLabel").grid(
					row=row, column=col_base, sticky="w", pady=2, padx=(0, 6)
				)
				ttk.Entry(form, textvariable=var, width=30).grid(
					row=row, column=col_base + 1, sticky="ew", pady=2, padx=(0, 12)
				)

			btns = ttk.Frame(form_wrap, style="Card.TFrame")
			btns.pack(fill="x", pady=(8, 0))
			ttk.Button(btns, text="Confirm Add Product", command=self._add, style="Primary.TButton").pack(side="left", padx=(0, 6))
			ttk.Button(btns, text="Update", command=self._update).pack(side="left", padx=(0, 6))
			ttk.Button(btns, text="Delete", command=self._delete, style="Danger.TButton").pack(side="left", padx=(0, 6))
			ttk.Button(btns, text="Reset Form", command=self._reset_form).pack(side="left")

		table_wrap = ttk.Frame(self, style="Card.TFrame")
		table_wrap.pack(fill="both", expand=True, pady=(8, 0))

		cols = ("product_id", "product_name", "category", "price", "quantity", "barcode", "added_by")
		self._tree = ttk.Treeview(table_wrap, columns=cols, show="headings", height=10, selectmode="browse")
		headers = {
			"product_id": ("ID", 60, "center"),
			"product_name": ("Name", 230, "w"),
			"category": ("Category", 130, "w"),
			"price": ("Price", 100, "center"),
			"quantity": ("Qty", 80, "center"),
			"barcode": ("Barcode", 160, "w"),
			"added_by": ("Added By", 100, "w"),
		}
		for col, (label, width, anchor) in headers.items():
			self._tree.heading(col, text=label)
			self._tree.column(col, width=width, anchor=anchor)
		self._tree.pack(expand=True, fill="both")
		self._tree.bind("<<TreeviewSelect>>", lambda _e: self._load_selected())

		ttk.Label(self, textvariable=self._status_var, style="InfoBadge.TLabel").pack(anchor="w", pady=(8, 0))

	def _clear_search(self) -> None:
		self._search_var.set("")
		self.refresh()

	def _reset_form(self) -> None:
		self._name_var.set("")
		self._category_var.set("")
		self._price_var.set("")
		self._quantity_var.set("")
		self._barcode_var.set("")
		self._tree.selection_remove(*self._tree.selection())

	def _selected_product_id(self) -> int | None:
		selected = self._tree.selection()
		if not selected:
			return None
		return int(self._tree.item(selected[0], "values")[0])

	def _validate_form(self) -> bool:
		if not is_non_empty(self._name_var.get().strip()):
			self._status_var.set("Name is required.")
			return False
		if not is_valid_price(self._price_var.get().strip()):
			self._status_var.set("Enter a valid price.")
			return False
		if not is_valid_quantity(self._quantity_var.get().strip()):
			self._status_var.set("Enter a valid quantity.")
			return False
		return True

	def _load_selected(self) -> None:
		selected = self._tree.selection()
		if not selected:
			return
		vals = self._tree.item(selected[0], "values")
		self._name_var.set(str(vals[1]))
		self._category_var.set(str(vals[2]))
		self._price_var.set(str(vals[3]))
		self._quantity_var.set(str(vals[4]))
		self._barcode_var.set(str(vals[5]))

	def _add(self) -> None:
		if not self._validate_form():
			return
		if not _show_confirmation_dialog(self, "Confirm Add", "Add this new product?"):
			return
		success = add_product(
			self._name_var.get().strip(),
			self._category_var.get().strip(),
			float(self._price_var.get().strip()),
			int(self._quantity_var.get().strip()),
			self._barcode_var.get().strip(),
			actor=self._user,
		)
		self._status_var.set("Product added." if success else "Add failed. Check barcode uniqueness.")
		if success:
			self._reset_form()
			self.refresh()

	def _update(self) -> None:
		pid = self._selected_product_id()
		if pid is None:
			self._status_var.set("Select a product row to update.")
			return
		if not self._validate_form():
			return
		if not _show_confirmation_dialog(self, "Confirm Update", "Update this product record?"):
			return
		success = update_product(
			pid,
			self._name_var.get().strip(),
			self._category_var.get().strip(),
			float(self._price_var.get().strip()),
			int(self._quantity_var.get().strip()),
			self._barcode_var.get().strip(),
		)
		self._status_var.set("Product updated." if success else "Update failed.")
		if success:
			self.refresh()

	def _delete(self) -> None:
		pid = self._selected_product_id()
		if pid is None:
			self._status_var.set("Select a product row to delete.")
			return
		if not _show_confirmation_dialog(self, "Delete Product", "Delete selected product?"):
			return
		success = delete_product(pid)
		self._status_var.set("Product deleted." if success else "Delete failed.")
		if success:
			self._reset_form()
			self.refresh()

	def refresh(self) -> None:
		for row in self._tree.get_children():
			self._tree.delete(row)
		query = self._search_var.get().strip()
		rows = search_products(query) if query else get_all_products()
		for p in rows:
			self._tree.insert(
				"",
				"end",
				values=(
					p["product_id"],
					p["product_name"],
					p["category"],
					f"{float(p['price']):.2f}",
					p["quantity"],
					p["barcode"],
					p.get("username", "N/A"),
				),
			)


class InventoryPanel(ttk.Frame):
	"""In-dashboard inventory panel with stock adjustment and logs."""

	def __init__(self, parent: tk.Misc, user: dict[str, Any] | None = None) -> None:
		super().__init__(parent, style="Card.TFrame", padding=(12, 10))
		self._user = user or {}
		self._role = str(self._user.get("role", "Cashier"))
		self._is_admin = self._role == "Admin"
		self._status_var = tk.StringVar()
		self._product_var = tk.StringVar()
		self._change_var = tk.StringVar()
		self._reason_var = tk.StringVar()
		self._search_var = tk.StringVar()
		self._build_ui()
		self.refresh()

	def _build_ui(self) -> None:
		top = ttk.Frame(self, style="Card.TFrame")
		top.pack(fill="x", pady=(0, 8))
		ttk.Label(top, text="Inventory", style="SectionTitle.TLabel").pack(side="left")
		ttk.Label(top, text="Search:", style="SectionSub.TLabel").pack(side="left", padx=(14, 6))
		search_entry = ttk.Entry(top, textvariable=self._search_var, width=30)
		search_entry.pack(side="left")
		search_entry.bind("<KeyRelease>", lambda _e: self.refresh())
		ttk.Button(top, text="Refresh", command=self.refresh, style="Primary.TButton").pack(side="right")

		nb = ttk.Notebook(self)
		nb.pack(expand=True, fill="both")

		stock_tab = ttk.Frame(nb, padding=6)
		log_tab = ttk.Frame(nb, padding=6)
		nb.add(stock_tab, text="Stock")
		nb.add(log_tab, text="Adjustments")

		self._stock_tree = ttk.Treeview(stock_tab, columns=("id", "name", "qty", "price", "barcode"), show="headings", height=10)
		for col, label, width, anchor in [
			("id", "ID", 60, "center"),
			("name", "Product", 260, "w"),
			("qty", "Qty", 80, "center"),
			("price", "Price", 100, "e"),
			("barcode", "Barcode", 180, "w"),
		]:
			self._stock_tree.heading(col, text=label)
			self._stock_tree.column(col, width=width, anchor=anchor)
		self._stock_tree.pack(expand=True, fill="both", pady=(0, 8))

		if not self._is_admin:
			form = ttk.Frame(stock_tab, style="Card.TFrame")
			form.pack(fill="x")
			ttk.Label(form, text="Product:", style="SectionSub.TLabel").grid(row=0, column=0, sticky="w", padx=(0, 6))
			self._product_combo = ttk.Combobox(form, textvariable=self._product_var, state="readonly", width=30)
			self._product_combo.grid(row=0, column=1, sticky="w")
			ttk.Label(form, text="Change:", style="SectionSub.TLabel").grid(row=0, column=2, sticky="w", padx=(12, 6))
			ttk.Entry(form, textvariable=self._change_var, width=10).grid(row=0, column=3, sticky="w")
			ttk.Label(form, text="Reason:", style="SectionSub.TLabel").grid(row=0, column=4, sticky="w", padx=(12, 6))
			ttk.Entry(form, textvariable=self._reason_var, width=24).grid(row=0, column=5, sticky="w")
			ttk.Button(form, text="Apply", command=self._apply_adjustment, style="Primary.TButton").grid(row=0, column=6, padx=(12, 0))

		self._log_tree = ttk.Treeview(log_tab, columns=("log_id", "product", "change", "reason", "date", "user"), show="headings", height=14)
		for col, label, width, anchor in [
			("log_id", "Log #", 70, "center"),
			("product", "Product", 220, "w"),
			("change", "Change", 90, "center"),
			("reason", "Reason", 280, "w"),
			("date", "Date", 170, "w"),
			("user", "User", 120, "w"),
		]:
			self._log_tree.heading(col, text=label)
			self._log_tree.column(col, width=width, anchor=anchor)
		self._log_tree.pack(expand=True, fill="both")

		ttk.Label(self, textvariable=self._status_var, style="InfoBadge.TLabel").pack(anchor="w", pady=(8, 0))

	def _apply_adjustment(self) -> None:
		selection = self._product_var.get().strip()
		if "ID:" not in selection:
			self._status_var.set("Select a product first.")
			return
		try:
			product_id = int(selection.split("ID:")[1].rstrip(")"))
			change = int(self._change_var.get().strip())
		except ValueError:
			self._status_var.set("Change must be a whole number.")
			return
		reason = self._reason_var.get().strip()
		if not reason:
			self._status_var.set("Reason is required.")
			return
		success = adjust_stock(product_id, change, reason, actor=self._user)
		self._status_var.set("Stock updated." if success else "Adjustment failed.")
		if success:
			self._change_var.set("")
			self._reason_var.set("")
			self.refresh()

	def refresh(self) -> None:
		for row in self._stock_tree.get_children():
			self._stock_tree.delete(row)
		products = get_all_products()
		query = self._search_var.get().strip().lower()
		if query:
			products = [
				p
				for p in products
				if query in str(p.get("product_id", "")).lower()
				or query in str(p.get("product_name", "")).lower()
				or query in str(p.get("barcode", "")).lower()
				or query in str(p.get("category", "")).lower()
			]
		combo_values: list[str] = []
		for p in products:
			self._stock_tree.insert(
				"",
				"end",
				values=(p["product_id"], p["product_name"], p["quantity"], f"{float(p['price']):.2f}", p["barcode"]),
			)
			combo_values.append(f"{p['product_name']} (ID:{p['product_id']})")
		
		if not self._is_admin:
			self._product_combo.configure(values=combo_values)
			if combo_values and not self._product_var.get():
				self._product_combo.current(0)

		for row in self._log_tree.get_children():
			self._log_tree.delete(row)
		for item in get_inventory_log():
			chg = int(item["change_amount"])
			self._log_tree.insert(
				"",
				"end",
				values=(
					item["inventory_id"],
					item["product_name"],
					f"{chg:+d}",
					item["reason"],
					item["date"],
					item.get("username", "N/A"),
				),
			)


class CustomersPanel(ttk.Frame):
	"""In-dashboard customer search, edit, and purchase history panel."""

	def __init__(self, parent: tk.Misc) -> None:
		super().__init__(parent, style="Card.TFrame", padding=(12, 10))
		self._search_var = tk.StringVar()
		self._name_var = tk.StringVar()
		self._phone_var = tk.StringVar()
		self._email_var = tk.StringVar()
		self._address_var = tk.StringVar()
		self._status_var = tk.StringVar()
		self._cust_rows: list[dict[str, Any]] = []
		self._build_ui()
		self.refresh()

	def _build_ui(self) -> None:
		head = ttk.Frame(self, style="Card.TFrame")
		head.pack(fill="x", pady=(0, 8))
		ttk.Label(head, text="Customers", style="SectionTitle.TLabel").pack(side="left")
		ttk.Button(head, text="Export CSV", command=self._export_csv).pack(side="right")

		search = ttk.Frame(self, style="Card.TFrame")
		search.pack(fill="x", pady=(0, 8))
		ttk.Label(search, text="Search:", style="SectionSub.TLabel").pack(side="left", padx=(0, 6))
		entry = ttk.Entry(search, textvariable=self._search_var, width=32)
		entry.pack(side="left")
		entry.bind("<KeyRelease>", lambda _e: self.refresh())
		ttk.Button(search, text="Clear", command=self._clear_search).pack(side="left", padx=(6, 0))

		nb = ttk.Notebook(self)
		nb.pack(expand=True, fill="both")

		main_tab = ttk.Frame(nb, padding=6)
		history_tab = ttk.Frame(nb, padding=6)
		nb.add(main_tab, text="Customers")
		nb.add(history_tab, text="Purchase History")

		form = ttk.Frame(main_tab, style="Card.TFrame")
		form.pack(fill="x")
		form.columnconfigure(1, weight=1)
		form.columnconfigure(3, weight=1)
		ttk.Label(form, text="Name:", style="SectionSub.TLabel").grid(row=0, column=0, sticky="w", pady=2, padx=(0, 6))
		ttk.Entry(form, textvariable=self._name_var).grid(row=0, column=1, sticky="ew", pady=2, padx=(0, 12))
		ttk.Label(form, text="Phone:", style="SectionSub.TLabel").grid(row=0, column=2, sticky="w", pady=2, padx=(0, 6))
		ttk.Entry(form, textvariable=self._phone_var).grid(row=0, column=3, sticky="ew", pady=2)

		ttk.Label(form, text="Email:", style="SectionSub.TLabel").grid(row=1, column=0, sticky="w", pady=2, padx=(0, 6))
		ttk.Entry(form, textvariable=self._email_var).grid(row=1, column=1, columnspan=3, sticky="ew", pady=2)
		ttk.Label(form, text="Address:", style="SectionSub.TLabel").grid(row=2, column=0, sticky="w", pady=2, padx=(0, 6))
		ttk.Entry(form, textvariable=self._address_var).grid(row=2, column=1, columnspan=3, sticky="ew", pady=2)

		btns = ttk.Frame(main_tab, style="Card.TFrame")
		btns.pack(fill="x", pady=(8, 0))
		ttk.Button(btns, text="Confirm Add Customer", command=self._add_customer, style="Primary.TButton").pack(side="left", padx=(0, 6))
		ttk.Button(btns, text="Update", command=self._update_customer).pack(side="left", padx=(0, 6))
		ttk.Button(btns, text="Reset Form", command=self._reset_form).pack(side="left")

		cols = ("customer_id", "name", "phone", "email", "address", "loyalty_points")
		self._tree = ttk.Treeview(main_tab, columns=cols, show="headings", height=10, selectmode="browse")
		for col, label, width, anchor in [
			("customer_id", "ID", 60, "center"),
			("name", "Name", 180, "w"),
			("phone", "Phone", 130, "w"),
			("email", "Email", 180, "w"),
			("address", "Address", 180, "w"),
			("loyalty_points", "Points", 70, "center"),
		]:
			self._tree.heading(col, text=label)
			self._tree.column(col, width=width, anchor=anchor)
		self._tree.pack(expand=True, fill="both", pady=(8, 0))
		self._tree.bind("<<TreeviewSelect>>", lambda _e: self._on_customer_selected())

		self._history_tree = ttk.Treeview(history_tab, columns=("sale_id", "date", "total_amount", "payment_method"), show="headings", height=12)
		for col, label, width, anchor in [
			("sale_id", "Sale #", 90, "center"),
			("date", "Date", 180, "w"),
			("total_amount", "Total", 110, "e"),
			("payment_method", "Method", 130, "w"),
		]:
			self._history_tree.heading(col, text=label)
			self._history_tree.column(col, width=width, anchor=anchor)
		self._history_tree.pack(expand=True, fill="both")

		ttk.Label(self, textvariable=self._status_var, style="InfoBadge.TLabel").pack(anchor="w", pady=(8, 0))

	def _clear_search(self) -> None:
		self._search_var.set("")
		self.refresh()

	def _selected_customer_id(self) -> int | None:
		selected = self._tree.selection()
		if not selected:
			return None
		return int(self._tree.item(selected[0], "values")[0])

	def _reset_form(self) -> None:
		self._name_var.set("")
		self._phone_var.set("")
		self._email_var.set("")
		self._address_var.set("")

	def _on_customer_selected(self) -> None:
		selected = self._tree.selection()
		if not selected:
			return
		vals = self._tree.item(selected[0], "values")
		self._name_var.set(str(vals[1]))
		self._phone_var.set(str(vals[2]))
		self._email_var.set(str(vals[3]))
		self._address_var.set(str(vals[4]))
		customer_id = int(vals[0])
		self._load_history(customer_id)

	def _load_history(self, customer_id: int) -> None:
		for row in self._history_tree.get_children():
			self._history_tree.delete(row)
		for sale in get_purchase_history(customer_id):
			self._history_tree.insert(
				"",
				"end",
				values=(sale["sale_id"], sale["date"], f"{float(sale['total_amount']):.2f}", sale["payment_method"]),
			)

	def _validate(self) -> bool:
		if not is_non_empty(self._name_var.get().strip()):
			self._status_var.set("Customer name is required.")
			return False
		email = self._email_var.get().strip()
		if email and not is_valid_email(email):
			self._status_var.set("Email format is invalid.")
			return False
		return True

	def _add_customer(self) -> None:
		if not self._validate():
			return
		if not _show_confirmation_dialog(self, "Confirm Add", "Add this customer record?"):
			return
		new_id = add_customer(
			self._name_var.get().strip(),
			self._phone_var.get().strip(),
			self._email_var.get().strip(),
			self._address_var.get().strip(),
		)
		if new_id:
			self._status_var.set("Customer added.")
			self._reset_form()
			self.refresh()
		else:
			self._status_var.set("Failed to add customer.")

	def _update_customer(self) -> None:
		cid = self._selected_customer_id()
		if cid is None:
			self._status_var.set("Select a customer row to update.")
			return
		if not self._validate():
			return
		if not _show_confirmation_dialog(self, "Confirm Update", "Update this customer record?"):
			return
		success = update_customer(
			cid,
			self._name_var.get().strip(),
			self._phone_var.get().strip(),
			self._email_var.get().strip(),
			self._address_var.get().strip(),
		)
		self._status_var.set("Customer updated." if success else "Update failed.")
		if success:
			self.refresh()

	def refresh(self) -> None:
		for row in self._tree.get_children():
			self._tree.delete(row)
		query = self._search_var.get().strip()
		rows = search_customers(query) if query else get_all_customers()
		self._cust_rows = list(rows)
		for c in rows:
			self._tree.insert(
				"",
				"end",
				values=(
					c["customer_id"],
					c["name"],
					c["phone"],
					c["email"],
					c["address"],
					c["loyalty_points"],
				),
			)

	def _export_csv(self) -> None:
		if not self._cust_rows:
			messagebox.showinfo("Export", "No customer data to export.", parent=self)
			return
		path = filedialog.asksaveasfilename(
			parent=self,
			defaultextension=".csv",
			filetypes=[("CSV Files", "*.csv")],
			initialfile="customers.csv",
		)
		if not path:
			return
		try:
			with open(path, "w", newline="", encoding="utf-8") as csv_file:
				writer = csv.DictWriter(
					csv_file,
					fieldnames=["customer_id", "name", "phone", "email", "address", "loyalty_points"],
				)
				writer.writeheader()
				for row in self._cust_rows:
					writer.writerow({k: row.get(k, "") for k in ["customer_id", "name", "phone", "email", "address", "loyalty_points"]})
			messagebox.showinfo("Export", "Customers exported successfully.", parent=self)
		except OSError as exc:
			messagebox.showerror("Export Error", str(exc), parent=self)


class UserManagementPanel(ttk.Frame):
	"""Admin panel for managing application users."""

	ROLES = ["Admin", "Manager", "Cashier"]
	GENDERS = ["M", "F", "Other"]

	def __init__(self, parent: tk.Misc, current_user: dict[str, Any] | None = None) -> None:
		super().__init__(parent, style="Card.TFrame", padding=(12, 10))
		self._current_user = current_user or {}
		self._fullname_var = tk.StringVar()
		self._address_var = tk.StringVar()
		self._gender_var = tk.StringVar(value="M")
		self._role_var = tk.StringVar(value="Cashier")
		self._status_var = tk.StringVar()
		self._username_display_var = tk.StringVar()
		self._password_display_var = tk.StringVar()
		self._build_ui()
		self.refresh()

	def _build_ui(self) -> None:
		header = ttk.Frame(self, style="Card.TFrame")
		header.pack(fill="x", pady=(0, 8))
		ttk.Label(header, text="User Management", style="SectionTitle.TLabel").pack(side="left")
		ttk.Button(header, text="Refresh", command=self.refresh, style="Primary.TButton").pack(side="right")

		# Form placed ABOVE the tree so it is always visible
		form = ttk.Frame(self, style="Card.TFrame")
		form.pack(fill="x", pady=(0, 8))
		form.columnconfigure(1, weight=1)
		form.columnconfigure(3, weight=1)

		# Row 0: Full Name | Gender
		ttk.Label(form, text="Full Name:", style="SectionSub.TLabel").grid(row=0, column=0, sticky="w", padx=(0, 6), pady=2)
		ttk.Entry(form, textvariable=self._fullname_var).grid(row=0, column=1, sticky="ew", pady=2, padx=(0, 12))
		ttk.Label(form, text="Gender:", style="SectionSub.TLabel").grid(row=0, column=2, sticky="w", padx=(0, 6), pady=2)
		gender_combo = ttk.Combobox(form, textvariable=self._gender_var, values=self.GENDERS, state="readonly", width=14)
		gender_combo.grid(row=0, column=3, sticky="ew", pady=2)

		# Row 1: Address | Role
		ttk.Label(form, text="Address:", style="SectionSub.TLabel").grid(row=1, column=0, sticky="w", padx=(0, 6), pady=2)
		ttk.Entry(form, textvariable=self._address_var).grid(row=1, column=1, sticky="ew", pady=2, padx=(0, 12))
		ttk.Label(form, text="Role:", style="SectionSub.TLabel").grid(row=1, column=2, sticky="w", padx=(0, 6), pady=2)
		role_combo = ttk.Combobox(form, textvariable=self._role_var, values=self.ROLES, state="readonly", width=14)
		role_combo.grid(row=1, column=3, sticky="ew", pady=2)

		# Auto-generated credentials display (read-only)
		creds_frame = ttk.LabelFrame(self, text="Auto-Generated Credentials", style="Card.TFrame", padding=(8, 6))
		creds_frame.pack(fill="x", pady=(0, 8))
		creds_frame.columnconfigure(1, weight=1)
		creds_frame.columnconfigure(3, weight=1)

		ttk.Label(creds_frame, text="Username:", style="SectionSub.TLabel").grid(row=0, column=0, sticky="w", padx=(0, 6), pady=2)
		ttk.Entry(creds_frame, textvariable=self._username_display_var, state="readonly").grid(row=0, column=1, sticky="ew", pady=2, padx=(0, 12))
		ttk.Label(creds_frame, text="Password:", style="SectionSub.TLabel").grid(row=0, column=2, sticky="w", padx=(0, 6), pady=2)
		ttk.Entry(creds_frame, textvariable=self._password_display_var, state="readonly").grid(row=0, column=3, sticky="ew", pady=2)

		# Buttons
		btns = ttk.Frame(self, style="Card.TFrame")
		btns.pack(fill="x", pady=(0, 8))
		ttk.Button(btns, text="Generate Credentials", command=self._generate_credentials, style="Primary.TButton").pack(side="left", padx=(0, 6))
		ttk.Button(btns, text="Add User", command=self._add_user, style="Primary.TButton").pack(side="left", padx=(0, 6))
		ttk.Button(btns, text="Update Role", command=self._update_role).pack(side="left", padx=(0, 6))
		ttk.Button(btns, text="Update User Details", command=self._update_user_details, style="Primary.TButton").pack(side="left", padx=(0, 6))
		ttk.Button(btns, text="Delete User", command=self._delete_user, style="Danger.TButton").pack(side="left")

		ttk.Label(self, textvariable=self._status_var, style="InfoBadge.TLabel").pack(anchor="w", pady=(0, 6))

		# Tree below the form — takes all remaining vertical space
		self._tree = ttk.Treeview(
			self,
			columns=("user_id", "username", "full_name", "address", "role"),
			show="headings",
			selectmode="browse",
			height=10,
		)
		for col, label, width, anchor in [
			("user_id", "ID", 50, "center"),
			("username", "Username", 140, "w"),
			("full_name", "Full Name", 180, "w"),
			("address", "Address", 220, "w"),
			("role", "Role", 100, "w"),
		]:
			self._tree.heading(col, text=label)
			self._tree.column(col, width=width, anchor=anchor)
		self._tree.pack(fill="both", expand=True)
		self._tree.bind("<<TreeviewSelect>>", lambda _e: self._on_select())

	def _selected_user_id(self) -> int | None:
		selected = self._tree.selection()
		if not selected:
			return None
		return int(self._tree.item(selected[0], "values")[0])

	def _selected_username(self) -> str:
		selected = self._tree.selection()
		if not selected:
			return ""
		return str(self._tree.item(selected[0], "values")[1])

	def _on_select(self) -> None:
		selected = self._tree.selection()
		if not selected:
			return
		vals = self._tree.item(selected[0], "values")
		self._fullname_var.set(str(vals[2]) if len(vals) > 2 else "")
		self._address_var.set(str(vals[3]) if len(vals) > 3 else "")
		self._role_var.set(str(vals[4]) if len(vals) > 4 else "")
		# Clear generated credentials when selecting
		self._username_display_var.set("")
		self._password_display_var.set("")

	def _generate_credentials(self) -> None:
		"""Generate username from full name and a secure password."""
		full_name = self._fullname_var.get().strip()
		if not full_name:
			self._status_var.set("Enter a full name first.")
			return
		
		# Generate username from full name
		username = generate_username_from_fullname(full_name)
		if not username:
			self._status_var.set("Invalid full name format.")
			return
		
		# Generate secure password
		password = generate_password(14)
		
		# Display generated credentials
		self._username_display_var.set(username)
		self._password_display_var.set(password)
		self._status_var.set("Credentials generated. Username blends all name parts and is capped to 7 letters.")

	def _add_user(self) -> None:
		if not has_permission(self._current_user, "manage_users"):
			self._status_var.set("Permission denied.")
			return
		
		full_name = self._fullname_var.get().strip()
		address = self._address_var.get().strip()
		gender = self._gender_var.get().strip()
		role = self._role_var.get().strip()
		username = self._username_display_var.get().strip()
		password = self._password_display_var.get().strip()
		
		if not full_name:
			self._status_var.set("Full name is required.")
			return
		if not username:
			self._status_var.set("Generate credentials first.")
			return
		if not password:
			self._status_var.set("Generate credentials first.")
			return
		if role not in self.ROLES:
			self._status_var.set("Select a valid role.")
			return
		
		if not _show_confirmation_dialog(self, "Create User", f"Create user '{username}' ({full_name}) with role {role}?"):
			return
		
		success = create_user(username, password, role, full_name=full_name, address=address, gender=gender, actor=self._current_user)
		if success:
			self._status_var.set("User added successfully.")
			self._download_credentials_docx(username, password, full_name, role)
			self._fullname_var.set("")
			self._address_var.set("")
			self._gender_var.set("M")
			self._username_display_var.set("")
			self._password_display_var.set("")
			self.refresh()
		else:
			self._status_var.set("Failed to add user.")

	def _download_credentials_docx(self, username: str, password: str, full_name: str, role: str) -> None:
		"""Save generated credentials as a DOCX file on desktop."""
		try:
			path = _save_credentials_docx_to_desktop(
				username=username,
				password=password,
				title="New User Credentials",
				full_name=full_name,
				role=role,
			)
			messagebox.showinfo(
				"Saved",
				f"Credentials saved to desktop:\n{path}",
				parent=self,
			)
		except Exception as exc:
			messagebox.showerror("Export Error", str(exc), parent=self)

	def _update_role(self) -> None:
		if not has_permission(self._current_user, "manage_users"):
			self._status_var.set("Permission denied.")
			return
		user_id = self._selected_user_id()
		role = self._role_var.get().strip()
		if user_id is None:
			self._status_var.set("Select a user first.")
			return
		if role not in self.ROLES:
			self._status_var.set("Select a valid role.")
			return
		if not _show_confirmation_dialog(self, "Update Role", "Update selected user role?"):
			return
		success = update_user_role(user_id, role, actor=self._current_user)
		self._status_var.set("Role updated." if success else "Role update failed.")
		if success:
			self.refresh()

	def _update_user_details(self) -> None:
		if not has_permission(self._current_user, "manage_users"):
			self._status_var.set("Permission denied.")
			return
		user_id = self._selected_user_id()
		if user_id is None:
			self._status_var.set("Select a user first.")
			return
		role = self._role_var.get().strip()
		if role not in self.ROLES:
			self._status_var.set("Select a valid role.")
			return
		gender = self._gender_var.get().strip() or "Other"
		if gender not in self.GENDERS:
			self._status_var.set("Select a valid gender.")
			return
		if not _show_confirmation_dialog(self, "Update User Details", "Update selected user's details and role?"):
			return
		success = update_user_details(
			user_id,
			self._fullname_var.get().strip(),
			self._address_var.get().strip(),
			gender,
			role,
			actor=self._current_user,
		)
		self._status_var.set("User details updated." if success else "User update failed.")
		if success:
			self.refresh()

	def _delete_user(self) -> None:
		if not has_permission(self._current_user, "manage_users"):
			self._status_var.set("Permission denied.")
			return
		user_id = self._selected_user_id()
		username = self._selected_username()
		if user_id is None:
			self._status_var.set("Select a user first.")
			return
		if username == str(self._current_user.get("username", "")):
			self._status_var.set("You cannot delete your own active account.")
			return
		if not _show_confirmation_dialog(self, "Delete User", f"Delete user '{username}'?"):
			return
		success = delete_user(user_id, actor=self._current_user)
		self._status_var.set("User deleted." if success else "Delete failed.")
		if success:
			self.refresh()

	def refresh(self) -> None:
		if not has_permission(self._current_user, "manage_users"):
			self._status_var.set("Permission denied.")
			return
		for row in self._tree.get_children():
			self._tree.delete(row)
		for user in get_all_users():
			self._tree.insert(
				"",
				"end",
				values=(
					user["user_id"],
					user["username"],
					user.get("full_name", ""),
					user.get("address", ""),
					user["role"],
				),
			)


class ChangePasswordPanel(ttk.Frame):
	"""Self-service password change panel for all roles."""

	def __init__(self, parent: tk.Misc, current_user: dict[str, Any] | None = None) -> None:
		super().__init__(parent, style="Card.TFrame", padding=(12, 10))
		self._current_user = current_user or {}
		self._old_password_var = tk.StringVar()
		self._old_password_confirm_var = tk.StringVar()
		self._new_password_var = tk.StringVar()
		self._new_password_confirm_var = tk.StringVar()
		self._show_passwords_var = tk.BooleanVar(value=False)
		self._status_var = tk.StringVar(value="Enter old password twice and new password twice.")
		self._build_ui()

	def _build_ui(self) -> None:
		header = ttk.Frame(self, style="Card.TFrame")
		header.pack(fill="x", pady=(0, 8))
		ttk.Label(header, text="Change Password", style="SectionTitle.TLabel").pack(side="left")

		form = ttk.Frame(self, style="Card.TFrame")
		form.pack(fill="x", pady=(0, 8))
		form.columnconfigure(1, weight=1)

		username = str(self._current_user.get("username", ""))
		ttk.Label(form, text="Username:", style="SectionSub.TLabel").grid(row=0, column=0, sticky="w", padx=(0, 8), pady=4)
		ttk.Entry(form, state="readonly", width=30, justify="left").grid(row=0, column=1, sticky="ew", pady=4)
		form.winfo_children()[-1].configure(state="normal")
		form.winfo_children()[-1].insert(0, username)
		form.winfo_children()[-1].configure(state="readonly")

		ttk.Label(form, text="Old Password:", style="SectionSub.TLabel").grid(row=1, column=0, sticky="w", padx=(0, 8), pady=4)
		self._old_password_entry = ttk.Entry(form, textvariable=self._old_password_var, show="*")
		self._old_password_entry.grid(row=1, column=1, sticky="ew", pady=4)

		ttk.Label(form, text="Confirm Old Password:", style="SectionSub.TLabel").grid(row=2, column=0, sticky="w", padx=(0, 8), pady=4)
		self._old_password_confirm_entry = ttk.Entry(form, textvariable=self._old_password_confirm_var, show="*")
		self._old_password_confirm_entry.grid(row=2, column=1, sticky="ew", pady=4)

		ttk.Label(form, text="New Password:", style="SectionSub.TLabel").grid(row=3, column=0, sticky="w", padx=(0, 8), pady=4)
		self._new_password_entry = ttk.Entry(form, textvariable=self._new_password_var, show="*")
		self._new_password_entry.grid(row=3, column=1, sticky="ew", pady=4)

		ttk.Label(form, text="Confirm New Password:", style="SectionSub.TLabel").grid(row=4, column=0, sticky="w", padx=(0, 8), pady=4)
		self._new_password_confirm_entry = ttk.Entry(form, textvariable=self._new_password_confirm_var, show="*")
		self._new_password_confirm_entry.grid(row=4, column=1, sticky="ew", pady=4)

		ttk.Checkbutton(
			form,
			text="Show passwords",
			variable=self._show_passwords_var,
			command=self._toggle_password_visibility,
		).grid(row=5, column=1, sticky="w", pady=(4, 0))

		buttons = ttk.Frame(self, style="Card.TFrame")
		buttons.pack(fill="x", pady=(0, 8))
		ttk.Button(buttons, text="Apply Password Change", command=self._apply_password_change, style="Primary.TButton").pack(side="left")

		ttk.Label(self, textvariable=self._status_var, style="InfoBadge.TLabel").pack(anchor="w", pady=(0, 8))

	def _apply_password_change(self) -> None:
		username = str(self._current_user.get("username", "")).strip()
		old_password = self._old_password_var.get().strip()
		old_password_confirm = self._old_password_confirm_var.get().strip()
		new_password = self._new_password_var.get().strip()
		new_password_confirm = self._new_password_confirm_var.get().strip()

		if not username:
			self._status_var.set("No active username found for this session.")
			return
		if not old_password or not old_password_confirm or not new_password or not new_password_confirm:
			self._status_var.set("All password fields are required.")
			return
		if old_password != old_password_confirm:
			self._status_var.set("Old password entries do not match.")
			return
		if new_password != new_password_confirm:
			self._status_var.set("New password entries do not match.")
			return
		error = password_policy_error(new_password)
		if error:
			self._status_var.set(error)
			return
		if "user_id" in self._current_user and is_password_in_use(new_password, exclude_user_id=int(self._current_user["user_id"])):
			self._status_var.set("That password is already used by another user. Choose a different password.")
			return

		success = change_password(username, old_password, new_password)
		if not success:
			self._status_var.set("Password change failed. Check your old password and try again.")
			return

		self._status_var.set("Password updated successfully.")
		self._old_password_var.set("")
		self._old_password_confirm_var.set("")
		self._new_password_var.set("")
		self._new_password_confirm_var.set("")

	def _toggle_password_visibility(self) -> None:
		show = "" if self._show_passwords_var.get() else "*"
		self._old_password_entry.configure(show=show)
		self._old_password_confirm_entry.configure(show=show)
		self._new_password_entry.configure(show=show)
		self._new_password_confirm_entry.configure(show=show)

class AuditLogsPanel(ttk.Frame):
	"""Admin panel to review audit/security event logs."""

	def __init__(self, parent: tk.Misc, current_user: dict[str, Any] | None = None) -> None:
		super().__init__(parent, style="Card.TFrame", padding=(12, 10))
		self._current_user = current_user or {}
		self._rows: list[dict[str, Any]] = []
		self._status_var = tk.StringVar(value="Ready")
		self._build_ui()
		self.refresh()

	def _build_ui(self) -> None:
		head = ttk.Frame(self, style="Card.TFrame")
		head.pack(fill="x", pady=(0, 8))
		ttk.Label(head, text="Audit Logs", style="SectionTitle.TLabel").pack(side="left")
		ttk.Button(head, text="Export CSV", command=self._export_csv).pack(side="right")
		ttk.Button(head, text="Refresh", command=self.refresh, style="Primary.TButton").pack(side="right", padx=(0, 8))

		cols = ("log_id", "timestamp", "username", "action", "detail", "integrity")
		self._tree = ttk.Treeview(self, columns=cols, show="headings", height=14)
		headers = {
			"log_id": ("ID", 70, "center"),
			"timestamp": ("Timestamp", 170, "w"),
			"username": ("User", 120, "w"),
			"action": ("Action", 140, "w"),
			"detail": ("Detail", 320, "w"),
			"integrity": ("Hash OK", 90, "center"),
		}
		for col, (label, width, anchor) in headers.items():
			self._tree.heading(col, text=label)
			self._tree.column(col, width=width, anchor=anchor)
		self._tree.pack(fill="both", expand=True)

		ttk.Label(self, textvariable=self._status_var, style="InfoBadge.TLabel").pack(anchor="w", pady=(8, 0))

	def refresh(self) -> None:
		if not has_permission(self._current_user, "view_audit_logs"):
			self._status_var.set("Permission denied.")
			return
		for row in self._tree.get_children():
			self._tree.delete(row)
		rows = list_recent_audit_logs(limit=1000, actor=self._current_user)
		self._rows = rows
		for row in rows:
			self._tree.insert(
				"",
				"end",
				values=(
					row["log_id"],
					row["timestamp"],
					row.get("username") or "-",
					row["action"],
					row.get("detail") or "",
					"Yes" if row.get("hash_ok") else "No",
				),
			)
		self._status_var.set(f"Loaded {len(rows)} audit events")

	def _export_csv(self) -> None:
		if not has_permission(self._current_user, "view_audit_logs"):
			self._status_var.set("Permission denied.")
			return
		if not self._rows:
			messagebox.showinfo("Export", "No audit logs to export.", parent=self)
			return
		path = filedialog.asksaveasfilename(
			parent=self,
			defaultextension=".csv",
			filetypes=[("CSV Files", "*.csv")],
			initialfile="audit_logs.csv",
		)
		if not path:
			return
		try:
			with open(path, "w", newline="", encoding="utf-8") as csv_file:
				writer = csv.DictWriter(
					csv_file,
					fieldnames=[
						"log_id",
						"timestamp",
						"user_id",
						"username",
						"action",
						"detail",
						"prev_hash",
						"row_hash",
						"hash_ok",
					],
				)
				writer.writeheader()
				for row in self._rows:
					writer.writerow(row)
			messagebox.showinfo("Export", "Audit logs exported successfully.", parent=self)
		except OSError as exc:
			messagebox.showerror("Export Error", str(exc), parent=self)


class ReportsPanel(ttk.Frame):
	"""In-dashboard reports panel with optional CSV export."""

	REPORT_TYPES = ["Daily Sales", "Weekly Sales", "Product Performance", "Inventory", "Cashier"]

	def __init__(self, parent: tk.Misc, user: dict[str, Any] | None = None) -> None:
		super().__init__(parent, style="Card.TFrame", padding=(12, 10))
		self._user = user or {}
		self._rows: list[dict[str, Any]] = []
		self._table_columns: list[str] = []
		self._report_var = tk.StringVar(value=self.REPORT_TYPES[0])
		today = datetime.now().strftime("%Y-%m-%d")
		self._start_var = tk.StringVar(value=today)
		self._end_var = tk.StringVar(value=today)
		self._cashier_var = tk.StringVar(value="")
		self._cashier_choices: list[str] = []
		self._status_var = tk.StringVar(value="Ready")
		self._build_ui()
		self._run_report()

	def _build_ui(self) -> None:
		controls = ttk.Frame(self, style="Card.TFrame")
		controls.pack(fill="x", pady=(0, 8))

		ttk.Label(controls, text="Report:", style="SectionSub.TLabel").grid(row=0, column=0, sticky="w", padx=(0, 6))
		combo = ttk.Combobox(controls, textvariable=self._report_var, values=self.REPORT_TYPES, state="readonly", width=20)
		combo.grid(row=0, column=1, sticky="w")
		combo.bind("<<ComboboxSelected>>", lambda _e: self._on_report_changed())

		ttk.Label(controls, text="Start:", style="SectionSub.TLabel").grid(row=0, column=2, sticky="w", padx=(12, 6))
		ttk.Entry(controls, textvariable=self._start_var, width=12).grid(row=0, column=3, sticky="w")
		ttk.Label(controls, text="End:", style="SectionSub.TLabel").grid(row=0, column=4, sticky="w", padx=(12, 6))
		ttk.Entry(controls, textvariable=self._end_var, width=12).grid(row=0, column=5, sticky="w")
		ttk.Label(controls, text="Cashier (ID/Username):", style="SectionSub.TLabel").grid(row=0, column=6, sticky="w", padx=(12, 6))
		self._cashier_entry = ttk.Combobox(controls, textvariable=self._cashier_var, width=20)
		self._cashier_entry.grid(row=0, column=7, sticky="w")
		self._refresh_cashier_choices()

		ttk.Button(controls, text="Run", command=self._run_report, style="Primary.TButton").grid(row=0, column=8, padx=(12, 6))
		ttk.Button(controls, text="Export CSV", command=self._export_csv).grid(row=0, column=9)

		table_wrap = ttk.Frame(self, style="Card.TFrame")
		table_wrap.pack(expand=True, fill="both")
		self._table = ttk.Treeview(table_wrap, columns=(), show="headings", selectmode="browse")
		vsb = ttk.Scrollbar(table_wrap, orient="vertical", command=self._table.yview)
		hsb = ttk.Scrollbar(table_wrap, orient="horizontal", command=self._table.xview)
		self._table.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
		self._table.pack(side="left", expand=True, fill="both")
		vsb.pack(side="right", fill="y")
		hsb.pack(side="bottom", fill="x")

		ttk.Label(self, textvariable=self._status_var, style="InfoBadge.TLabel").pack(anchor="w", pady=(8, 0))
		self._on_report_changed()

	def _on_report_changed(self) -> None:
		if self._report_var.get() == "Cashier":
			self._cashier_entry.config(state="readonly")
			self._refresh_cashier_choices()
		else:
			self._cashier_entry.config(state="disabled")

	def _refresh_cashier_choices(self) -> None:
		cashier_users = [
			u
			for u in get_all_users()
			if str(u.get("role", "")).strip() == "Cashier"
		]
		choices: list[str] = []
		for user in cashier_users:
			uid = str(user.get("user_id", "")).strip()
			uname = str(user.get("username", "")).strip()
			if uid and uname:
				choices.append(f"{uid} | {uname}")
			if uname:
				choices.append(uname)
		self._cashier_choices = sorted(set(choices), key=lambda v: v.lower())
		self._cashier_entry.configure(values=self._cashier_choices)
		current = self._cashier_var.get().strip()
		if self._cashier_choices and current not in self._cashier_choices:
			self._cashier_var.set(self._cashier_choices[0])

	def _run_report(self) -> None:
		report = self._report_var.get()
		start = self._start_var.get().strip()
		end = self._end_var.get().strip()
		try:
			if report == "Daily Sales":
				data = daily_sales_report(start)
				top = "; ".join(f"{p['product_name']}({p['units_sold']})" for p in data["top_products"]) or "-"
				rows = [{"date": data["date"], "total_sales": data["total_sales"], "transactions": data["transactions"], "top_products": top}]
			elif report == "Weekly Sales":
				rows = weekly_sales_report(start)
			elif report == "Product Performance":
				rows = product_performance_report()
			elif report == "Inventory":
				rows = inventory_report()
			else:
				cashier_ref = self._cashier_var.get().strip()
				if not cashier_ref:
					self._status_var.set("Enter cashier ID or username.")
					return
				if "|" in cashier_ref:
					cashier_ref = cashier_ref.split("|", 1)[0].strip()
				data = cashier_report(cashier_ref, start, end)
				rows = [{
					"cashier_id": data["user_id"],
					"username": data["username"],
					"role": data["role"],
					"start": data["start"],
					"end": data["end"],
					"total_sales": data["total_sales"],
					"transactions": data["transactions"],
				}]
		except Exception as exc:
			self._status_var.set(str(exc))
			return
		self._render_rows(rows)

	def _render_rows(self, rows: list[dict[str, Any]]) -> None:
		for item in self._table.get_children():
			self._table.delete(item)
		if not rows:
			self._table.configure(columns=())
			self._rows = []
			self._table_columns = []
			self._status_var.set("No data found")
			return
		columns = list(rows[0].keys())
		self._table.configure(columns=columns)
		for col in columns:
			self._table.heading(col, text=col.replace("_", " ").title())
			self._table.column(col, width=140, anchor="w")
		for row in rows:
			self._table.insert("", "end", values=[row.get(col, "") for col in columns])
		self._rows = rows
		self._table_columns = columns
		self._status_var.set(f"{len(rows)} row(s) displayed")

	def _export_csv(self) -> None:
		if not self._rows or not self._table_columns:
			self._status_var.set("Run a report before export.")
			return
		path = filedialog.asksaveasfilename(
			parent=self,
			defaultextension=".csv",
			filetypes=[("CSV Files", "*.csv")],
			initialfile="report_export.csv",
		)
		if not path:
			return
		try:
			with open(path, "w", newline="", encoding="utf-8") as csv_file:
				writer = csv.DictWriter(csv_file, fieldnames=self._table_columns)
				writer.writeheader()
				writer.writerows(self._rows)
			self._status_var.set("Report exported successfully.")
		except OSError as exc:
			self._status_var.set(f"Export failed: {exc}")


class PasswordResetTicketsPanel(ttk.Frame):
	"""Admin panel for managing password reset tickets."""

	def __init__(self, parent: tk.Misc, current_user: dict[str, Any] | None = None) -> None:
		super().__init__(parent, style="Card.TFrame", padding=(12, 10))
		self._current_user = current_user or {}
		self._status_var = tk.StringVar(value="Ready")
		self._new_password_var = tk.StringVar()
		self._status_filter_var = tk.StringVar(value="OPEN")
		self._build_ui()
		self.refresh()
		self.after(5000, self._auto_refresh)

	def _build_ui(self) -> None:
		header = ttk.Frame(self, style="Card.TFrame")
		header.pack(fill="x", pady=(0, 8))
		ttk.Label(header, text="Password Reset Tickets", style="SectionTitle.TLabel").pack(side="left")
		ttk.Label(header, text="Status:", style="SectionSub.TLabel").pack(side="left", padx=(12, 6))
		status_combo = ttk.Combobox(
			header,
			textvariable=self._status_filter_var,
			values=["OPEN", "RESOLVED", "CLOSED", "ALL"],
			state="readonly",
			width=12,
		)
		status_combo.pack(side="left")
		status_combo.bind("<<ComboboxSelected>>", lambda _e: self.refresh())
		ttk.Button(header, text="Refresh", command=self.refresh, style="Primary.TButton").pack(side="right")

		# Instructions
		ttk.Label(
			self,
			text="Manage password reset requests from users. Select a ticket and generate a temporary password.",
			style="CardSub.TLabel",
		).pack(anchor="w", pady=(0, 8))

		# Buttons
		btns = ttk.Frame(self, style="Card.TFrame")
		btns.pack(fill="x", pady=(0, 8))
		ttk.Button(btns, text="Generate Password", command=self._generate_password, style="Primary.TButton").pack(side="left", padx=(0, 6))
		ttk.Button(btns, text="Close Ticket", command=self._close_ticket).pack(side="left")

		ttk.Label(self, textvariable=self._status_var, style="InfoBadge.TLabel").pack(anchor="w", pady=(0, 6))

		# Treeview for tickets
		self._tree = ttk.Treeview(
			self,
			columns=("ticket_id", "username", "email", "created_at", "status"),
			show="headings",
			selectmode="browse",
			height=10,
		)
		for col, label, width, anchor in [
			("ticket_id", "ID", 60, "center"),
			("username", "Username", 150, "w"),
			("email", "Email", 200, "w"),
			("created_at", "Created", 150, "w"),
			("status", "Status", 100, "center"),
		]:
			self._tree.heading(col, text=label)
			self._tree.column(col, width=width, anchor=anchor)
		self._tree.pack(fill="both", expand=True)

	def _selected_ticket(self) -> dict[str, Any] | None:
		"""Get the selected ticket from the tree."""
		selected = self._tree.selection()
		if not selected:
			return None
		values = self._tree.item(selected[0], "values")
		return {
			"ticket_id": int(values[0]),
			"username": str(values[1]),
			"email": str(values[2]),
			"created_at": str(values[3]),
			"status": str(values[4]),
		}

	def _generate_password(self) -> None:
		"""Generate a temporary password and show confirmation dialog."""
		if not has_permission(self._current_user, "manage_users"):
			self._status_var.set("Permission denied.")
			return

		ticket = self._selected_ticket()
		if not ticket:
			self._status_var.set("Select a ticket first.")
			return

		if ticket["status"] != "OPEN":
			self._status_var.set(f"Cannot resolve a {ticket['status'].lower()} ticket.")
			return

		# Generate a temporary password
		temp_password = generate_password()

		# Show confirmation dialog with generated password
		dialog = tk.Toplevel(self)
		dialog.title("Generated Temporary Password")
		dialog.geometry("450x280")
		dialog.resizable(False, False)
		dialog.transient(self.winfo_toplevel())
		dialog.grab_set()
		apply_modern_theme(dialog)

		container = ttk.Frame(dialog, padding=20)
		container.pack(expand=True, fill="both")

		ttk.Label(container, text="Temporary Password Generated", style="CardTitle.TLabel").pack(anchor="w", pady=(0, 12))
		
		info_text = f"Username: {ticket['username']}\nTemporary Password:"
		ttk.Label(container, text=info_text, style="CardSub.TLabel").pack(anchor="w", pady=(0, 8))

		# Password display with copy-friendly formatting
		password_frame = ttk.Frame(container, relief="solid", borderwidth=1)
		password_frame.pack(fill="x", pady=(0, 16), ipady=8, ipadx=8)
		
		password_label = ttk.Label(password_frame, text=temp_password, font=("Courier", 12, "bold"), foreground="#1976d2")
		password_label.pack()

		ttk.Label(container, text="Share this password with the user securely.", style="CardSub.TLabel", foreground="#666").pack(anchor="w", pady=(0, 16))

		# Buttons
		button_frame = ttk.Frame(container)
		button_frame.pack(fill="x")

		def _confirm_and_resolve() -> None:
			"""Confirm and apply the generated password."""
			success = resolve_ticket(ticket["ticket_id"], temp_password, actor=self._current_user)
			self._status_var.set("Password reset successfully!" if success else "Failed to reset password.")
			if success:
				for item_id in self._tree.get_children():
					vals = self._tree.item(item_id, "values")
					if vals and int(vals[0]) == int(ticket["ticket_id"]):
						self._tree.item(item_id, values=(vals[0], vals[1], vals[2], vals[3], "RESOLVED"))
						break
				self._status_filter_var.set("ALL")
				try:
					path = _save_credentials_docx_to_desktop(
						username=ticket["username"],
						password=temp_password,
						title="Ticket Reset Credentials",
					)
					messagebox.showinfo(
						"Saved",
						f"Reset credentials saved to desktop:\n{path}",
						parent=self,
					)
				except Exception as exc:
					messagebox.showerror("Export Error", str(exc), parent=self)
				self.refresh()
			dialog.destroy()

		ttk.Button(button_frame, text="Confirm & Apply", command=_confirm_and_resolve, style="Primary.TButton").pack(side="left", padx=(0, 8))
		ttk.Button(button_frame, text="Cancel", command=dialog.destroy).pack(side="left")

		# Focus on the window
		dialog.focus_set()


	def _close_ticket(self) -> None:
		"""Close a ticket without resetting the password."""
		if not has_permission(self._current_user, "manage_users"):
			self._status_var.set("Permission denied.")
			return

		ticket = self._selected_ticket()
		if not ticket:
			self._status_var.set("Select a ticket first.")
			return

		if ticket["status"] != "OPEN":
			self._status_var.set(f"Cannot close a {ticket['status'].lower()} ticket.")
			return

		if not _show_confirmation_dialog(
			self,
			"Close Ticket",
			f"Close ticket for '{ticket['username']}'?\n\nNo password reset will be performed.",
		):
			return

		success = close_ticket(ticket["ticket_id"], actor=self._current_user)
		self._status_var.set("Ticket closed." if success else "Failed to close ticket.")
		if success:
			self.refresh()

	def refresh(self) -> None:
		"""Reload all tickets from the database."""
		if not has_permission(self._current_user, "manage_users"):
			self._status_var.set("Permission denied.")
			return

		for row in self._tree.get_children():
			self._tree.delete(row)

		tickets = list_all_tickets()
		status_filter = self._status_filter_var.get().strip().upper()
		if status_filter and status_filter != "ALL":
			tickets = [ticket for ticket in tickets if str(ticket.get("status", "")).upper() == status_filter]
		for ticket in tickets:
			self._tree.insert(
				"",
				"end",
				values=(
					ticket["ticket_id"],
					ticket["username"],
					ticket.get("email") or "-",
					ticket["created_at"],
					ticket["status"],
				),
			)

		self._status_var.set(f"Loaded {len(tickets)} ticket(s)")

	def _auto_refresh(self) -> None:
		try:
			self.refresh()
		except Exception:
			pass
		if self.winfo_exists():
			self.after(5000, self._auto_refresh)

